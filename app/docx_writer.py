from docx import Document
from docx.shared import Pt, RGBColor
import re

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(14)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        paragraph.add_run(text[cursor:start])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = end
    paragraph.add_run(text[cursor:])

def add_code_block(doc, code_lines):
    para = doc.add_paragraph()
    run = para.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(10)

def add_markdown_table(doc, lines):
    if len(lines) < 2:
        return
    headers = [cell.strip(" *") for cell in lines[0].split("|") if cell.strip()]
    rows = [
        [cell.strip() for cell in row.split("|") if cell.strip()]
        for row in lines[2:]  # skip header and separator
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell

def create_docx(ts_text: str, buffer):
    doc = Document()
    doc.add_heading('TECHNICAL SPECIFICATION', level=1)
    lines = ts_text.splitlines()
    current_section = ""
    current_content = []
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    section_header_pattern = re.compile(r"^\s*(\d{1,2})\.\s*(.+?)(:|$)")
    subheading_pattern = re.compile(r"^\s*(\d{1,2})\.(\d+)\s+(.+)")
    table_line_pattern = re.compile(r"^\|(.+?)\|$")

    def flush_current_content():
        if current_section:
            add_heading(doc, current_section)
        for para in current_content:
            add_paragraph(doc, para)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Handle code blocks (start/end)
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block and code_block_lines:
                add_code_block(doc, code_block_lines)
                code_block_lines = []
            continue
        elif in_code_block:
            code_block_lines.append(line)
            continue

        # Handle tables
        if table_line_pattern.match(line):
            table_lines.append(line)
            in_table = True
            continue
        elif in_table and table_lines:
            flush_current_content()
            current_content = []
            add_markdown_table(doc, table_lines)
            table_lines = []
            in_table = False
            continue

        # Section/header detection
        m_section = section_header_pattern.match(line)
        m_sub = subheading_pattern.match(line)
        if m_section:
            if current_section or current_content:
                flush_current_content()
                current_content = []
            current_section = m_section.group(0)
        elif m_sub:
            add_subheading(doc, m_sub.group(0))
        else:
            current_content.append(line)

    # Final flush
    if in_table and table_lines:
        flush_current_content()
        add_markdown_table(doc, table_lines)
    elif current_section or current_content:
        flush_current_content()
    doc.save(buffer)